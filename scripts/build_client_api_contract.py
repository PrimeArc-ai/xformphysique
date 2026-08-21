from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "xform-client-api-contract.docx"

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
MID_BLUE = "1F4D78"
INK = "1F2933"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_FILL = "F5F7FA"
WHITE = "FFFFFF"
GREEN = "2F6B3F"
CAUTION = "7A5A00"


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade(cell_or_paragraph, fill):
    target = cell_or_paragraph._tc if hasattr(cell_or_paragraph, "_tc") else cell_or_paragraph._p
    props = target.get_or_add_tcPr() if hasattr(target, "get_or_add_tcPr") else target.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table_pr = table._tbl.tblPr
    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = table_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_keep(paragraph, keep_with_next=False, keep_together=False):
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together


def add_text(doc, text, style="Normal", color=None, bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_font(first, color=color, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_font(rest, color=color)
    else:
        run = paragraph.add_run(text)
        set_font(run, color=color)
    return paragraph


def add_label(doc, text, tone=BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text.upper())
    set_font(run, size=8.5, color=tone, bold=True)
    set_keep(paragraph, keep_with_next=True)
    return paragraph


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shd)
    run = paragraph.add_run(text)
    set_font(run, name="Consolas", size=7.7, color="233044")
    return paragraph


def add_endpoint(doc, number, method, route, ui, purpose, request=None, response=None, extra=None):
    heading = doc.add_paragraph(style="Heading 2")
    set_keep(heading, keep_with_next=True)
    method_run = heading.add_run(f"{number}. {method}  ")
    set_font(method_run, color=GREEN if method == "GET" else MID_BLUE, bold=True)
    route_run = heading.add_run(route)
    set_font(route_run, name="Consolas", size=11.5, color=DARK_BLUE, bold=True)
    meta = doc.add_paragraph()
    set_keep(meta, keep_with_next=True)
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(4)
    ui_run = meta.add_run("UI: ")
    set_font(ui_run, size=9, color=MUTED, bold=True)
    value_run = meta.add_run(ui)
    set_font(value_run, size=9, color=MUTED)
    body = doc.add_paragraph()
    set_keep(body, keep_with_next=True)
    body.paragraph_format.space_after = Pt(4)
    run = body.add_run(purpose)
    set_font(run, size=10.5, color=INK)
    if request:
        add_label(doc, "Request")
        add_code(doc, request)
    if response:
        add_label(doc, "Response · 200")
        add_code(doc, response)
    if extra:
        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(0)
        note.paragraph_format.space_after = Pt(8)
        note.paragraph_format.left_indent = Inches(0.08)
        run = note.add_run(extra)
        set_font(run, size=9.5, color=CAUTION, italic=True)


def endpoint_data():
    return [
        ("1", "GET", "/dashboard", "Client Dashboard", "Loads client identity, current body signal, seven-day trend, check-in status, training volume and next actions.", "GET /api/v1/client/dashboard", '''{
  "client": {"id": "cl_001", "first_name": "Maya", "primary_goal": "body_recomposition"},
  "body": {"current_weight_kg": 68.4, "latest_waist_cm": 71.0, "change_from_start_kg": -1.7, "target_progress_percent": 74, "trend": [{"date": "2026-08-20", "weight_kg": 68.4}]},
  "check_ins": {"count": 1, "status": "submitted"},
  "training_volume": {"range_days": 30, "total_kg": 4280, "sessions": 12, "training_days": 8, "best_day_kg": 860, "daily_kg": [{"date": "2026-08-20", "volume_kg": 420}]},
  "next_actions": [{"type": "body_entry", "label": "Log body progress", "due": true}]
}''', None),
        ("2", "GET", "/body-entries?from=&to=&limit=", "Body Tracker", "Returns body-history records for chart, latest values and history table.", "GET /api/v1/client/body-entries?from=2026-08-01&to=2026-08-31&limit=60", '''{
  "items": [{"id": "body_001", "date": "2026-08-20", "weight_kg": 68.4, "waist_cm": 71.0, "created_at": "2026-08-20T08:30:00Z"}],
  "summary": {"start_weight_kg": 70.1, "latest_weight_kg": 68.4, "weight_change_kg": -1.7, "seven_day_average_kg": 68.8}
}''', None),
        ("3", "PUT", "/body-entries/{date}", "Body Tracker save", "Creates or replaces one daily body entry. Idempotent route prevents duplicate dates.", '''PUT /api/v1/client/body-entries/2026-08-20
Content-Type: application/json

{"weight_kg": 68.4, "waist_cm": 71.0}''', '''{
  "id": "body_001", "date": "2026-08-20", "weight_kg": 68.4, "waist_cm": 71.0,
  "summary": {"weight_change_kg": -1.7, "seven_day_average_kg": 68.8}
}''', "Reject future dates, non-positive weight and invalid units."),
        ("4", "GET", "/check-ins?limit=", "Check-ins, Health Summary", "Returns weekly check-in history plus client-specific due state.", "GET /api/v1/client/check-ins?limit=12", '''{
  "schedule": {"day_of_week": "sunday", "current_status": "due", "due_on": "2026-08-23"},
  "items": [{"id": "checkin_001", "period_start": "2026-08-17", "submitted_at": "2026-08-20T09:00:00Z", "energy_score": 4, "sleep_score": 3, "sentiment": "good", "observation": "Training felt consistent.", "concern": "Right knee felt sensitive after walking."}]
}''', None),
        ("5", "PUT", "/check-ins/current", "Weekly Check-in submit", "Creates or replaces current scheduled-cycle check-in.", '''PUT /api/v1/client/check-ins/current
Content-Type: application/json

{"energy_score": 4, "sleep_score": 3, "sentiment": "good", "observation": "Training and meals were consistent.", "concern": "Right knee felt sensitive after long walks."}''', '''{
  "id": "checkin_001", "period_start": "2026-08-17", "submitted_at": "2026-08-20T09:00:00Z", "status": "submitted",
  "energy_score": 4, "sleep_score": 3, "sentiment": "good", "observation": "Training and meals were consistent.", "concern": "Right knee felt sensitive after long walks."
}''', None),
        ("6", "GET", "/progress-photos?view=&limit=", "Progress Photos gallery and comparison", "Returns authorized photo metadata. UI filters by front, side or back and compares returned records locally.", "GET /api/v1/client/progress-photos?view=front&limit=50", '''{
  "items": [{"id": "photo_001", "view": "front", "captured_on": "2026-08-20", "file_name": "front.jpg", "content_url": "/api/v1/client/progress-photos/photo_001/content"}]
}''', None),
        ("7", "POST", "/progress-photos", "Progress Photos upload", "Stores one private client photo and creates its metadata record.", '''POST /api/v1/client/progress-photos
Content-Type: multipart/form-data

file=front.jpg
view=front
captured_on=2026-08-20''', '''{
  "id": "photo_001", "view": "front", "captured_on": "2026-08-20", "file_name": "front.jpg",
  "content_url": "/api/v1/client/progress-photos/photo_001/content"
}''', "Request is multipart because image binary is not JSON. Validate file type, size and ownership; never return public storage URLs."),
        ("8", "GET", "/progress-photos/{photo_id}/content", "Progress Photos image display", "Returns protected image bytes for one authorized photo.", "GET /api/v1/client/progress-photos/photo_001/content", None, "Response is image/jpeg, image/png or image/webp; no JSON response. Ownership check remains mandatory."),
        ("9", "GET", "/nutrition/active-plan?date=", "Nutrition page", "Returns only client-assigned plan, meal quantities, macros, restrictions and daily adherence.", "GET /api/v1/client/nutrition/active-plan?date=2026-08-20", '''{
  "plan_id": "mealplan_001", "name": "Recomposition baseline", "date": "2026-08-20",
  "daily_targets": {"calories_kcal": 1860, "protein_g": 135, "carbs_g": 190, "fat_g": 62},
  "restrictions": ["dairy_aware", "shellfish_free"],
  "meals": [{"id": "meal_001", "time": "08:00", "name": "Greek yoghurt bowl", "ingredients": [{"name": "Greek yoghurt", "quantity": 200, "unit": "g"}], "calories_kcal": 420, "macros": {"protein_g": 35, "carbs_g": 41, "fat_g": 13}, "adherence_status": "followed"}]
}''', "Client can read assigned plan only. Coach edits remain coach-only."),
        ("10", "PUT", "/nutrition/meals/{meal_id}/adherence", "Nutrition meal status", "Upserts followed, partly or missed status for one assigned meal and date.", '''PUT /api/v1/client/nutrition/meals/meal_001/adherence
Content-Type: application/json

{"date": "2026-08-20", "status": "followed"}''', '''{
  "meal_id": "meal_001", "date": "2026-08-20", "status": "followed",
  "daily_summary": {"logged_meals": 1, "total_meals": 3}
}''', "Allowed status: followed, partly, missed."),
        ("11", "POST", "/nutrition/recipe-guides", "Nutrition recipe helper", "Generates preparation guide from server-read assigned meal, ingredients and restrictions.", '''POST /api/v1/client/nutrition/recipe-guides
Content-Type: application/json

{"meal_id": "meal_001"}''', '''{
  "meal_id": "meal_001", "meal_name": "Greek yoghurt bowl",
  "guide": "Layer yoghurt, oats, berries and chia. Rest five minutes before serving.",
  "uses_assigned_ingredients_only": true, "remaining_requests_today": 1
}''', "Do not accept arbitrary ingredients or quantity changes. No diagnosis, supplements or unrestricted AI advice."),
        ("12", "GET", "/workout-sessions/today?date=", "Workout page", "Returns current assigned session, coach note, exercise prescriptions and completion state.", "GET /api/v1/client/workout-sessions/today?date=2026-08-20", '''{
  "session_id": "session_001", "date": "2026-08-20", "title": "Lower body strength", "week_label": "Week 03",
  "coach_note": "Move smoothly. Leave two reps in reserve.", "status": "ready", "estimated_duration_minutes": 45,
  "exercises": [{"plan_exercise_id": "pex_001", "order": 1, "name": "Goblet squat", "prescription": {"sets": 3, "reps": "10", "rest_seconds": 90}}]
}''', None),
        ("13", "PUT", "/workout-sessions/{session_id}", "Workout session log", "Saves overall difficulty, note and actual per-set reps/load; marks session completed when requested.", '''PUT /api/v1/client/workout-sessions/session_001
Content-Type: application/json

{
  "status": "completed", "completed_at": "2026-08-20T18:40:00Z", "overall_difficulty": "moderate", "note": "Good session. Knee felt comfortable.",
  "exercise_logs": [{"plan_exercise_id": "pex_001", "sets": [{"set_number": 1, "reps": 10, "load_kg": 20, "difficulty": "moderate"}]}]
}''', '''{
  "session_id": "session_001", "status": "completed", "completed_at": "2026-08-20T18:40:00Z", "volume_kg": 600, "completion_percent": 100
}''', None),
        ("14", "GET", "/health-summary", "Health Summary", "Read-only coach-approved planning context. No diagnosis or blood-report interpretation.", "GET /api/v1/client/health-summary", '''{
  "wellbeing": {"energy_score": 4, "sentiment": "good", "source_check_in_id": "checkin_001"},
  "planning_context": {"dietary_preferences": ["dairy_aware"], "allergies": ["shellfish"], "training_considerations": ["Monitor right knee comfort"], "coach_note": "Prioritise consistent sleep this week."},
  "safety_notice": "Coaching support only. Not medical advice."
}''', None),
        ("15", "GET", "/profile", "Profile page initial data", "Returns client goal, targets, timezone, schedule and planning preferences.", "GET /api/v1/client/profile", '''{
  "client_id": "cl_001", "name": "Maya Shah", "email": "maya@example.com", "primary_goal": "body_recomposition", "target_weight_kg": 65.0,
  "check_in_day": "sunday", "timezone": "Asia/Kolkata", "dietary_preferences": "Dairy-aware. Prefer quick weekday meals.",
  "allergies_injuries": "Shellfish allergy. Right knee sensitive after long walks."
}''', None),
        ("16", "PATCH", "/profile", "Profile page save", "Partially updates client-editable profile and planning preferences.", '''PATCH /api/v1/client/profile
Content-Type: application/json

{"primary_goal": "body_recomposition", "target_weight_kg": 65.0, "check_in_day": "sunday", "timezone": "Asia/Kolkata", "dietary_preferences": "Dairy-aware. Prefer quick weekday meals.", "allergies_injuries": "Shellfish allergy. Right knee sensitive after long walks."}''', '''{
  "client_id": "cl_001", "updated_at": "2026-08-20T10:30:00Z",
  "profile": {"primary_goal": "body_recomposition", "target_weight_kg": 65.0, "check_in_day": "sunday", "timezone": "Asia/Kolkata", "dietary_preferences": "Dairy-aware. Prefer quick weekday meals.", "allergies_injuries": "Shellfish allergy. Right knee sensitive after long walks."}
}''', None),
    ]


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (("Heading 1", 16, BLUE, 14, 7), ("Heading 2", 13, BLUE, 10, 5), ("Heading 3", 12, MID_BLUE, 8, 4)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("XFORM COACHING OS  |  CLIENT API CONTRACT")
    set_font(run, size=8, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("CLIENT API CONTRACT  ·  ")
    set_font(run, size=8, color=MUTED)
    add_page_number(footer)

    doc.core_properties.title = "XForm Coaching OS — Client API Contract"
    doc.core_properties.subject = "Frontend-to-FastAPI API contract for client pages"
    doc.core_properties.author = "XForm Coaching OS"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(20)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("TECHNICAL REFERENCE · FRONTEND CONTRACT")
    set_font(run, size=9, color=BLUE, bold=True)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("XForm Coaching OS\nClient API Contract")
    set_font(run, size=25, color=DARK_BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Client pages first · FastAPI + SQLite implementation target · Draft v0.1")
    set_font(run, size=11, color=MUTED)

    metadata = doc.add_table(rows=3, cols=2)
    set_table_geometry(metadata, [1800, 7560])
    metadata.style = "Table Grid"
    for row, label, value in zip(metadata.rows, ["Scope", "Base path", "Identity rule"], ["Client Dashboard, Body Tracker, Check-ins, Photos, Nutrition, Workout, Health Summary and Profile", "/api/v1/client", "Future access token determines client identity; client ID never travels in client routes"]):
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        label_run = row.cells[0].paragraphs[0].add_run(label)
        set_font(label_run, size=9, color=DARK_BLUE, bold=True)
        value_run = row.cells[1].paragraphs[0].add_run(value)
        set_font(value_run, size=9.5, color=INK)

    doc.add_paragraph()
    doc.add_paragraph("API rules", style="Heading 1")
    rules = doc.add_table(rows=1, cols=1)
    set_table_geometry(rules, [9360])
    set_cell_shading(rules.cell(0, 0), LIGHT_GRAY)
    para = rules.cell(0, 0).paragraphs[0]
    text = "Dates use YYYY-MM-DD. All writes return saved server data. Standard errors use 401 for no session, 403 for ownership/role failure, 404 for missing records, 409 for state conflict and 422 for validation failure. Health and recipe boundaries remain non-diagnostic."
    run = para.add_run(text)
    set_font(run, size=9.5, color=INK)
    add_label(doc, "Common error JSON", tone=CAUTION)
    add_code(doc, '''{
  "error": {"code": "validation_error", "message": "Weight must be greater than zero", "fields": {"weight_kg": "Must be greater than zero"}},
  "request_id": "req_..."
}''')

    doc.add_paragraph("Endpoint overview", style="Heading 1")
    add_text(doc, "Six client-facing domains and sixteen documented routes. The contracts below are grouped by the page or action they support, so the frontend and FastAPI implementation can be built in the same order.", color=INK)

    groups = [
        ("Dashboard and body data", [0, 1, 2]),
        ("Weekly check-ins", [3, 4]),
        ("Progress photos", [5, 6, 7]),
        ("Nutrition", [8, 9, 10]),
        ("Workout logging", [11, 12]),
        ("Health context and profile", [13, 14, 15]),
    ]
    data = endpoint_data()
    for group_title, indexes in groups:
        doc.add_paragraph(group_title, style="Heading 1")
        for endpoint_index in indexes:
            # Keep the complete profile-editing pair together on the closing page.
            # This avoids a response block starting a new page without its route context.
            if endpoint_index == 14:
                doc.add_page_break()
            add_endpoint(doc, *data[endpoint_index])

    doc.add_paragraph("Deferred routes", style="Heading 1")
    add_text(doc, "Authentication, coach plan creation, target editing, notifications, exports, audit events, consent retention and account deletion are outside this first client-API pass.", color=INK)
    add_label(doc, "IMPLEMENTATION NOTE")
    add_text(doc, "SQLite fits the first pass. Keep validation, role checks, client ownership, calculations, recipe restrictions and private-file access inside FastAPI; React should call these contracts and update its local cache from returned records.", color=INK)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
